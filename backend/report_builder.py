"""
Monthly report builder — generates .docx matching reference format.

Reference format (from user-provided .docx):
  - Title: "Resources Usage Metric Report" (Arial 24pt bold)
  - Subtitle: "Report Period: <Month Year>" (Arial 14pt)
  - Table (2 cols): "Item" | "Usage Metric"
    - Per VM: "Virtual Machine" | agent alias (bold)
    - Per metric: metric display name | line chart PNG

Uses ``python-docx`` + ``matplotlib`` (Agg backend).
"""

from __future__ import annotations

import logging
import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

# ── matplotlib MUST use Agg backend first ──────────────────────────
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt            # noqa: E402
import matplotlib.dates as mdates          # noqa: E402
import matplotlib.ticker as mticker        # noqa: E402

from docx import Document                  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402

logger = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────
CHART_DPI = 120
CHART_W_INCHES = 4.8
CHART_H_INCHES = 1.9

COLOR_CPU = "#0D6EFD"
COLOR_MEM = "#198754"
COLOR_DISK = "#DC3545"

FONT_FAMILY = "DejaVu Sans"

# Keywords for classifying modules by name (case-insensitive)
_CPU_NAMES = {"cpu load", "cpu usage", "cpu utilization", "cpu user", "cpu system"}
_MEM_NAMES = {"mem", "memory", "ram", "swap"}
_DISK_NAMES = {"disk", "storage", "drive", "volume", "c:", "d:", "e:", "f:", "/", "hdd", "ssd"}
# Skip these non-metric modules
_SKIP_NAMES = {
    "host alive", "host_live", "latency", "ping", "snmp", "status",
    "service", "process", "tcp", "udp", "network", "bandwidth",
    "uptime", "connection", "error", "event", "alert", "unknown",
    "traffic", "packet", "interface", "port", "agent",
    "sensor", "temperature", "humidity", "power",
}


def _classify_module(mod: dict) -> tuple[str, str, str]:
    """Classify by module name (from AJAX API).

    Returns ``(category, display_name, color)``.
    """
    name = (mod.get("module_name") or "").strip()
    lower = name.lower()

    # Skip list
    for skip_word in _SKIP_NAMES:
        if skip_word in lower:
            return ("skip", name, COLOR_DISK)

    # CPU
    for kw in _CPU_NAMES:
        if kw in lower:
            return ("cpu", name, COLOR_CPU)

    # Memory
    for kw in _MEM_NAMES:
        if kw in lower:
            return ("memory", name, COLOR_MEM)

    # Disk
    for kw in _DISK_NAMES:
        if kw in lower:
            return ("disk", name, COLOR_DISK)

    # Unknown name — fall back to value range
    points = mod.get("data_points", [])
    vals = [p["value"] for p in points if p.get("value") is not None]
    if vals:
        max_val = max(vals)
        if max_val <= 150:
            return ("cpu", name or "CPU", COLOR_CPU)
        avg = sum(vals) / len(vals)
        if avg > 500:
            return ("memory", name or "Memory", COLOR_MEM)
    return ("disk", name or "Disk Usage", COLOR_DISK)


def _should_show(category: str) -> bool:
    """Only CPU, Memory, and Disk metrics pass."""
    return category in ("cpu", "memory", "disk")


# ── Chart generation ──────────────────────────────────────────────────────

def _setup_style() -> None:
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": [FONT_FAMILY, "Liberation Sans", "Arial"],
        "font.size": 8,
        "axes.titlesize": 10,
        "axes.labelsize": 8,
        "xtick.labelsize": 7,
        "ytick.labelsize": 7,
        "figure.facecolor": "white",
        "axes.facecolor": "#FAFBFC",
        "axes.edgecolor": "#DEE2E6",
        "axes.grid": True,
        "grid.alpha": 0.35,
        "grid.color": "#CED4DA",
        "axes.spines.top": False,
        "axes.spines.right": False,
    })


_setup_style()


def generate_metric_chart(
    data_points: list[dict[str, Any]],
    label: str,
    color: str,
    output_path: str | Path,
) -> str | None:
    """Generate a line chart from pre-parsed data points.

    Args:
        data_points: List of ``{timestamp: datetime, value: float}`` dicts.
        label: Y-axis label.
        color: Line color.
        output_path: Where to save the PNG.

    Returns:
        Absolute path to the saved PNG, or None if no valid data.
    """
    if not data_points:
        return None

    # Sort by timestamp
    sorted_pts = sorted(data_points, key=lambda p: p.get("timestamp", datetime.min))

    timestamps = [p["timestamp"] for p in sorted_pts]
    values = [p["value"] for p in sorted_pts]

    if not timestamps or len(timestamps) < 2:
        return None

    fig, ax = plt.subplots(figsize=(CHART_W_INCHES, CHART_H_INCHES))

    ax.plot(timestamps, values, color=color, linewidth=1.2, marker=None)
    ax.fill_between(timestamps, values, alpha=0.08, color=color)

    # Y-axis label
    ax.set_ylabel(label, color=color, fontsize=7, fontweight="bold")

    # X-axis date formatting
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d %b"))
    locator = mdates.AutoDateLocator(minticks=3, maxticks=6)
    ax.xaxis.set_major_locator(locator)

    # Y-axis
    ax.yaxis.set_major_locator(mticker.MaxNLocator(nbins=4))

    # Average line
    avg = sum(values) / len(values)
    ax.axhline(y=avg, color=color, linestyle="--", linewidth=0.7, alpha=0.5)
    ax.text(timestamps[-1], avg, f"  avg {avg:.1f}",
            fontsize=6, color=color, va="center", alpha=0.8)

    fig.autofmt_xdate()
    fig.tight_layout(pad=0.5)
    fig.savefig(str(output_path), dpi=CHART_DPI, bbox_inches="tight")
    plt.close(fig)
    return str(Path(output_path).resolve())


# ── Document builder ──────────────────────────────────────────────────────

class ReportBuilder:
    """Builds the Resources Usage Metric Report .docx."""

    def __init__(
        self,
        tenant_name: str,
        period: str,
        output_dir: Path,
        date_start: str = "",
        date_end: str = "",
    ) -> None:
        self.tenant_name = tenant_name
        self.period = period
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        # Date range filter (YYYY-MM-DD)
        self._ts_start = _to_ts(date_start, end_of_day=False) if date_start else None
        self._ts_end = _to_ts(date_end, end_of_day=True) if date_end else None

        self.doc = Document()
        self._chart_paths: list[str] = []
        self._vm_count = 0

        # Page setup — portrait A4
        section = self.doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        # Default style
        style = self.doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

        self._add_header()

    # ── Header ────────────────────────────────────────────────────────

    def _add_header(self) -> None:
        """Title + subtitle + tenant name."""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Resources Usage Metric Report")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = "Arial"

        self.doc.add_paragraph("")

        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"Report Period: {self.period}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.name = "Arial"

        tn = self.doc.add_paragraph()
        tn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tn.add_run(self.tenant_name)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
        run.font.name = "Arial"

        self.doc.add_paragraph("")

        # Create table
        self._table = self.doc.add_table(rows=1, cols=2)
        self._table.style = "Table Grid"

        hdr = self._table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Usage Metric"
        for ci in (0, 1):
            for p in hdr[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(12)
                    r.font.bold = True
                    r.font.name = "Arial"

    # ── VM block ──────────────────────────────────────────────────────

    def add_vm_block(
        self,
        agent: dict,
        modules: list[dict],
        temp_dir: str,
    ) -> None:
        """Add one VM row + metric chart rows to the table.

        Args:
            agent: Agent dict from PandoraClient.
            modules: List from ``discover_agent_modules()``, each with:
                ``module_id``, ``data_points``, ``avg``, ``max_val``.
            temp_dir: Directory for temporary chart PNGs.
        """
        agent_alias = agent.get("alias", "Unknown")
        agent_name = agent.get("name", "")

        # ── VM header row ──────────────────────────────────────────
        vm_row = self._table.add_row()
        vm_label = vm_row.cells[0].paragraphs[0]
        run = vm_label.add_run("Virtual Machine")
        run.font.size = Pt(12)
        run.font.name = "Arial"

        vm_value = vm_row.cells[1].paragraphs[0]
        run = vm_value.add_run(agent_alias)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = "Arial"
        if agent_name:
            extra = vm_value.add_run(f"  [{agent_name}]")
            extra.font.size = Pt(9)

        self._vm_count += 1

        # Sort modules by ID (preserves Pandora's creation order)
        modules_sorted = sorted(modules, key=lambda m: m.get("module_id", 0))

        if not modules_sorted:
            self._add_no_data_row("CPU")
            self._add_no_data_row("Memory")
            self._add_no_data_row("Disk")
            return

        # Classify each module → keep only CPU, Memory, Disk
        cpu_mods: list[dict] = []
        mem_mods: list[dict] = []
        disk_mods: list[dict] = []

        for mod in modules_sorted:
            cat, _, _ = _classify_module(mod)
            if cat == "cpu":
                cpu_mods.append(mod)
            elif cat == "memory":
                mem_mods.append(mod)
            elif cat == "disk":
                disk_mods.append(mod)
            # skip: non-metric modules (Host Alive, ping, etc.)

        # Display: all CPU + all Memory + all Disk (no forced slots)
        to_display: list[tuple[str, str, dict]] = []

        for m in cpu_mods:
            to_display.append((m.get("module_name", "CPU"), COLOR_CPU, m))
        for m in mem_mods:
            to_display.append((m.get("module_name", "Memory"), COLOR_MEM, m))
        for m in disk_mods:
            to_display.append((m.get("module_name", "Disk"), COLOR_DISK, m))

        for label, color, mod in to_display:
            all_points = mod.get("data_points", [])

            # Filter to selected month
            in_month = []
            if all_points and self._ts_start is not None and self._ts_end is not None:
                in_month = _filter_by_month(all_points, self._ts_start, self._ts_end)
            elif all_points:
                in_month = all_points

            row = self._table.add_row()

            # Label cell
            label_cell = row.cells[0].paragraphs[0]
            run = label_cell.add_run(label)
            run.font.size = Pt(12)
            run.font.name = "Arial"

            # Chart cell
            if in_month and len(in_month) >= 2:
                chart_path = Path(temp_dir) / f"_chart_{agent.get('id_agente','?')}_{mod.get('module_id','?')}.png"
                result = generate_metric_chart(in_month, label, color, chart_path)
                if result:
                    self._chart_paths.append(result)
                    chart_para = row.cells[1].paragraphs[0]
                    run = chart_para.add_run()
                    run.add_picture(str(chart_path), width=Inches(CHART_W_INCHES))
                else:
                    self._add_empty_cell(row.cells[1], "Chart render failed")
            elif in_month and len(in_month) == 1:
                p = in_month[0]
                val_str = f"{p['value']:.2f}"
                ts_str = p['timestamp'].strftime("%d %b %Y %H:%M")
                self._add_empty_cell(row.cells[1], f"Single point: {val_str}  ({ts_str})")
            elif not mod:
                self._add_empty_cell(row.cells[1], "No data for this period")
            else:
                self._add_empty_cell(row.cells[1], "No data for this period")

        # Set column widths (apply to all rows)
        for row_obj in self._table.rows:
            row_obj.cells[0].width = Inches(1.8)
            row_obj.cells[1].width = Inches(6.0)

    def _add_no_data_row(self, label: str) -> None:
        """Add a row showing 'No data' for a metric."""
        row = self._table.add_row()
        run = row.cells[0].paragraphs[0].add_run(label)
        run.font.size = Pt(12)
        run.font.name = "Arial"
        self._add_empty_cell(row.cells[1], "No data")

    def _add_empty_cell(self, cell, text: str) -> None:
        """Fill a cell with placeholder italic text."""
        p = cell.paragraphs[0]
        # Clear default empty paragraph content
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xAD, 0xB5, 0xBD)

    # ── Save ──────────────────────────────────────────────────────────

    def save(self, filename_stem: str) -> str:
        """Save document to output_dir and return path."""
        safe = _safe_filename(filename_stem)
        path = self.output_dir / f"{safe}.docx"
        i = 1
        while path.exists():
            path = self.output_dir / f"{safe}_{i}.docx"
            i += 1

        self.doc.save(str(path))
        logger.info("Saved report: %s (%d VMs)", path, self._vm_count)

        for cp in self._chart_paths:
            try:
                os.unlink(cp)
            except OSError:
                pass

        return str(path.resolve())


# ── Top-level convenience ──────────────────────────────────────────────────

def build_report(
    tenant_name: str,
    period: str,
    agents: list[dict],
    agent_modules_map: dict[int, list[dict]],
    output_dir: str | Path,
    date_start: str = "",
    date_end: str = "",
) -> str:
    """Generate the Resources Usage Metric Report.

    Args:
        tenant_name: E.g. "PT Asuransi Central Asia [ACA]".
        period: E.g. "Mei 2026".
        agents: List of agent dicts.
        agent_modules_map: ``{agent_id: [module_dict, ...]}`` per agent.
        output_dir: Where to save the .docx.
        date_start: Start date "YYYY-MM-DD" (filters chart data).
        date_end: End date "YYYY-MM-DD" (filters chart data).

    Returns:
        Absolute path to the generated file.
    """
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as tmpdir:
        builder = ReportBuilder(
            tenant_name=tenant_name,
            period=period,
            output_dir=out,
            date_start=date_start,
            date_end=date_end,
        )
        for agent in agents:
            aid_raw = agent.get("id_agente")
            if aid_raw is None:
                continue
            aid = int(aid_raw)
            mods = agent_modules_map.get(aid, [])
            builder.add_vm_block(agent, mods, tmpdir)

        safe_tenant = _safe_filename(tenant_name)
        safe_period = period.replace(" ", "_")
        return builder.save(f"{safe_tenant}_{safe_period}")


def _safe_filename(name: str) -> str:
    name = name.strip().replace(" ", "_")
    name = re.sub(r"[^\w\-_.\[\]]", "", name)
    return name or "report"


def _to_ts(date_str: str, end_of_day: bool = False) -> int:
    """Convert YYYY-MM-DD to Unix timestamp."""
    dt = datetime.strptime(date_str.strip(), "%Y-%m-%d")
    if end_of_day:
        dt = dt.replace(hour=23, minute=59, second=59)
    return int(dt.timestamp())


def _filter_by_month(
    data_points: list[dict[str, Any]],
    ts_start: int,
    ts_end: int,
) -> list[dict[str, Any]]:
    """Filter data points to only those within the date range."""
    return [
        p for p in data_points
        if ts_start <= p.get("utimestamp", 0) <= ts_end
    ]
